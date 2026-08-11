from copy import deepcopy
from hashlib import sha256
from pathlib import Path
from shutil import copy2
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = next(Path('/Users/hunchulchoi/Downloads').glob('*IT_*.docx'))
OUTPUT = Path('/Users/hunchulchoi/projects/workspace/dgst/output/documents/푸른IT_스킬인벤토리_최훈철_완성본.docx')
CANDIDATE = Path('/Users/hunchulchoi/projects/workspace/dgst/tmp/skill-inventory/candidate.docx')
EXPECTED_SHA = '38e4ef8592f84c203114d70dd06667c17196165017125b21294157abcf928095'


def digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def clear_runs(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_cell(cell, text, *, size=8, bold=False, align='center'):
    """Replace a template value while retaining the cell and paragraph geometry."""
    p = cell.paragraphs[0]
    clear_runs(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Arial Unicode MS'
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:ascii'), 'Arial Unicode MS')
    rfonts.set(qn('w:hAnsi'), 'Arial Unicode MS')
    rfonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    p.alignment = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
    }[align]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


if digest(SOURCE) != EXPECTED_SHA:
    raise SystemExit('Reference document changed; fresh template distillation required.')

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
copy2(SOURCE, CANDIDATE)
doc = Document(CANDIDATE)

# Personal information
t = doc.tables[1]
set_cell(t.cell(0, 1), '최훈철', size=10, bold=True)
set_cell(t.cell(0, 6), '1978년', size=9)
set_cell(t.cell(0, 9), '남', size=9)
set_cell(t.cell(1, 1), '푸른IT', size=9)
set_cell(t.cell(1, 8), '20년 6개월', size=9)
set_cell(t.cell(2, 1), '개발팀', size=9)
set_cell(t.cell(2, 4), '차장', size=9)
set_cell(t.cell(2, 7), '1998.10~2000.12', size=8)
set_cell(t.cell(2, 9), '병장 / 고급', size=8)
set_cell(t.cell(3, 2), '', size=8)
set_cell(t.cell(3, 7), '', size=8)

# Contact details
t = doc.tables[2]
set_cell(t.cell(0, 1), '010-8660-5032', size=9, align='left')
set_cell(t.cell(1, 1), 'choihc165@gmail.com', size=9, align='left')
set_cell(t.cell(2, 1), '', size=9, align='left')
set_cell(t.cell(3, 1), '제주특별자치도 제주시 애월읍 곽지리', size=9, align='left')

# Education and certifications
t = doc.tables[3]
set_cell(t.cell(0, 0), '한림공업고등학교 졸업  1997년', size=8, align='left')
set_cell(t.cell(1, 0), '제주산업정보대학 전산학과 졸업  2002년 6월', size=8, align='left')
set_cell(t.cell(2, 0), '', size=8)
set_cell(t.cell(3, 0), '', size=8)
certs = [
    ('코딩지도사 1급', '2018.05'),
    ('정보처리산업기사', '2004.12'),
    ('SCJP', '2004.08'),
]
for row, (name, date) in enumerate(certs, start=1):
    set_cell(t.cell(row, 2), name, size=8, align='left')
    set_cell(t.cell(row, 3), date, size=8)

# Recent career summary
t = doc.tables[4]
career = [
    ('㈜비와이아이엔씨(BY inc.)', '2021.08~2026.04', '개발 차장', '공무원연금공단 맞춤형복지 개발·운영 / 레거시 현대화·AI 챗봇'),
    ('(주)퀀텀솔루션', '2021.03~2021.08', '책임연구원', 'Spring/JPA/Hadoop 서버 Docker 전환 및 GS마크 획득'),
    ('에스디시스템', '2019.02~2021.02', '개발 차장', '제주 C-ITS 연계 개발 / Python 문서 자동화'),
    ('이지팜', '2018.09~2019.01', '개발 차장', '검역본부 축산물 안전관리 시스템 고도화'),
]
for row, values in enumerate(career, start=1):
    for col, value in enumerate(values):
        set_cell(t.cell(row, col), value, size=7.2, align='left' if col in (0, 3) else 'center')

# Training and representative skill levels
t = doc.tables[5]
training = [
    ('(사)KH정보교육원', '2004.06', '2004.12', '자바 엔터프라이즈 전문가 과정'),
    ('JDC', '2019.08', '2019.08', '블록체인 구조와 프로그래밍'),
]
for row, values in enumerate(training, start=1):
    for col, value in enumerate(values):
        set_cell(t.cell(row, col), value, size=7.2, align='left' if col in (0, 3) else 'center')
skills = [('Java', '상'), ('Spring', '상'), ('SQL', '상'), ('Python', '중'), ('JavaScript', '중')]
for row, (name, level) in enumerate(skills, start=1):
    set_cell(t.cell(row, 5), name, size=8)
    set_cell(t.cell(row, 6), level, size=8, bold=True)

# Ten project/assignment records for the landscape skill inventory.
t = doc.tables[7]
projects = [
    ('맞춤형복지 시스템 고도화', '21.08~26.04', '공무원연금공단', 'BY inc.', '공공', 'Web/AI', '개발·운영', '서버', '', 'Java/Python/JS', '', 'Spring/TDD', 'JWE', 'LangChain/Pinecone'),
    ('Linux 서버 Docker 재구축', '21.03~21.08', '', '퀀텀솔루션', '솔루션', '서버', '재구축', '서버', 'Linux', 'Java', '', 'Spring Boot/JPA', '', 'Hadoop/Docker'),
    ('제주 C-ITS 실증사업', '19.02~21.02', '제주 C-ITS', '에스디시스템', 'ITS', '연계/자동화', '개발', '서버', '', 'Java/Python', '', 'Spring Boot', '', 'python-docx/openpyxl'),
    ('축산물 안전관리 고도화', '18.09~19.01', '검역본부', '이지팜', '공공', 'Web', '개발', '서버', '', 'Java/JS', '', '전자정부F/W', '', ''),
    ('맞춤형복지 개발·운영', '16.01~18.05', '공무원연금공단', '삼흥시스템', '공공', 'Web', '개발·운영', '서버', '', 'Java/JS', 'Tibero 5', '전자정부3.1', '', 'eXria'),
    ('SpringEnglish 서비스', '15.02~18.05', 'SpringEnglish', 'SpringEnglish', '교육', 'Web/결제', '팀장', 'EC2', 'Ubuntu', 'Java/JS', 'MariaDB', '전자정부F/W', 'Google+/PayPal', ''),
    ('인력포털 차세대 구축', '15.09~15.12', '한국관광공사', '아사달', '관광', '포털', '개발', '서버', '', '', '', '', '', ''),
    ('뉴스수신·내부정보시스템', '15.05~15.07', '광주유니버시아드', '광주U대회', '행사', '뉴스/업무', '개발', '서버', '', 'Java/JS/PHP', '', 'Netty/WordPress', '', 'jQuery'),
    ('입시 웹사이트 고도화', '15.02~15.05', '두원공대', '두원공대', '교육', 'Web', '팀장', '서버', '', 'Java/JS', 'MS-SQL', 'Spring/MyBatis', '아이핀', 'jQuery'),
    ('모바일 중고장터·가격비교', '14.11~15.02', '교보문고', '교보문고', '유통', '모바일', '개발', '서버', '', 'Java/JS', '', 'Spring/jQuery', '결제연동', ''),
]
for row, values in enumerate(projects, start=2):
    for col, value in enumerate(values):
        align = 'left' if col in (0, 4, 5, 6, 11, 13) else 'center'
        set_cell(t.cell(row, col), value, size=6.3, align=align)

# Fill the author text box. It is duplicated in DrawingML/VML fallback markup,
# so update every matching text-box run rather than only the visible branch.
for node in doc._element.xpath('.//w:txbxContent//w:t'):
    if node.text == '작성자':
        next_run = node.getparent().getnext()
        if next_run is not None:
            values = next_run.xpath('./w:t')
            if values:
                values[0].text = ' : 최훈철'

# The template names 굴림체, which is not present in this environment and
# otherwise renders Hangul as blanks. Arial Unicode MS is an available Korean
# sans-serif substitute; retain all source sizes, weights and geometry. Apply
# it to every body run, including DrawingML/VML text-box copies.
for run in doc._element.xpath('.//w:r'):
    rpr = run.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:ascii'), 'Arial Unicode MS')
    rfonts.set(qn('w:hAnsi'), 'Arial Unicode MS')
    rfonts.set(qn('w:eastAsia'), 'Arial Unicode MS')

doc.save(CANDIDATE)

# Transplant only the edited document part into the retained package. This
# avoids python-docx's harmless-but-broad rewrites of styles, settings,
# numbering and relationships, keeping every preserve-only part byte-identical.
with ZipFile(CANDIDATE, 'r') as candidate_zip:
    document_xml = candidate_zip.read('word/document.xml')
building = OUTPUT.with_name(OUTPUT.stem + '.building.docx')
with ZipFile(SOURCE, 'r') as source_zip, ZipFile(building, 'w', ZIP_DEFLATED) as output_zip:
    for info in source_zip.infolist():
        data = document_xml if info.filename == 'word/document.xml' else source_zip.read(info.filename)
        output_zip.writestr(info, data)
building.replace(OUTPUT)
print(OUTPUT)
