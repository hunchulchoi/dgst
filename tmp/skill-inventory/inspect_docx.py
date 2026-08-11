from pathlib import Path
from docx import Document

path = next(Path('/Users/hunchulchoi/Downloads').glob('*IT_*.docx'))
doc = Document(path)
print(f'paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}')
for pi, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'P{pi}: {p.text!r}')
for ti, table in enumerate(doc.tables):
    print(f'\nTABLE {ti}: rows={len(table.rows)} cols={len(table.columns)}')
    for ri, row in enumerate(table.rows):
        vals = []
        for ci, cell in enumerate(row.cells):
            text = ' / '.join(p.text.replace('\n', ' ') for p in cell.paragraphs).strip()
            vals.append(f'C{ci}[{id(cell._tc)}]={text!r}')
        print(f'R{ri}: ' + ' | '.join(vals))
