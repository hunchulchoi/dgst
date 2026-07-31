from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SANS = "Noto Sans CJK KR"
SERIF = "Noto Serif CJK KR"
BLACK = "000000"
LIGHT_GRAY = "E7E6E6"


def set_run_font(run, name: str, size: float, bold: bool | None = None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), name)


def set_paragraph_border(paragraph, *, bottom=None):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    if bottom:
        edge = OxmlElement("w:bottom")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(bottom.get("size", 8)))
        edge.set(qn("w:space"), str(bottom.get("space", 1)))
        edge.set(qn("w:color"), bottom.get("color", BLACK))
        pbdr.append(edge)


def set_table_borders(table, color=BLACK, size=8, inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    names = ["top", "left", "bottom", "right"]
    if inside:
        names += ["insideH", "insideV"]
    for name in names:
        tag = qn(f"w:{name}")
        old = borders.find(tag)
        if old is not None:
            borders.remove(old)
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)
        borders.append(edge)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "nil")


def set_table_geometry(table, widths_twips: list[int], total_twips: int):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_twips):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, *, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def clear_cell(cell):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    return paragraph


def fill_cell(
    cell,
    text: str,
    *,
    size=8.5,
    bold=False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    font=SANS,
):
    paragraph = clear_cell(cell)
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    return paragraph


def add_label_value(paragraph, label: str, value: str):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.05
    label_run = paragraph.add_run(label)
    set_run_font(label_run, SERIF, 9.7, True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, SERIF, 9.7, False)


def add_body_paragraph(document, text: str):
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.58)
    p.paragraph_format.first_line_indent = Cm(-0.58)
    p.paragraph_format.right_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.28
    run = p.add_run(text)
    set_run_font(run, SERIF, 10.2)
    return p


def build(output_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = SERIF
    normal.font.size = Pt(10.2)
    normal._element.rPr.rFonts.set(qn("w:ascii"), SERIF)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), SERIF)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), SERIF)
    normal.paragraph_format.space_after = Pt(0)

    list_style = doc.styles["List Number"]
    list_style.font.name = SERIF
    list_style.font.size = Pt(10.2)
    list_style._element.rPr.rFonts.set(qn("w:ascii"), SERIF)
    list_style._element.rPr.rFonts.set(qn("w:hAnsi"), SERIF)
    list_style._element.rPr.rFonts.set(qn("w:eastAsia"), SERIF)

    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company.paragraph_format.space_after = Pt(1)
    run = company.add_run("㈜타미우스골프앤빌리지")
    set_run_font(run, SANS, 18, True)

    address = doc.add_paragraph()
    address.alignment = WD_ALIGN_PARAGRAPH.CENTER
    address.paragraph_format.space_after = Pt(2)
    run = address.add_run(
        "63040 제주특별자치도 제주시 애월읍 화전길 201 / "
        "TEL(064)793-0703 / FAX (064)793-0709"
    )
    set_run_font(run, SANS, 8.0, True)
    set_paragraph_border(
        address, bottom={"size": 12, "space": 3, "color": "7F7F7F"}
    )

    outer = doc.add_table(rows=1, cols=2)
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(outer, [4400, 5240], 9640)
    remove_table_borders(outer)

    left = outer.cell(0, 0)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(left, top=110, start=0, bottom=0, end=110)
    p = clear_cell(left)
    add_label_value(p, "문서번호 : ", "타미우스-2026-01")
    add_label_value(left.add_paragraph(), "일    자 : ", "2026. 07. 30.")
    add_label_value(left.add_paragraph(), "수    신 : ", "KCTV 제주방송")
    add_label_value(left.add_paragraph(), "참    조 : ", "네트워크 / 보안 담당자")

    right = outer.cell(0, 1)
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(right, top=75, start=0, bottom=0, end=0)
    clear_cell(right)
    approval = right.add_table(rows=5, cols=5)
    approval.alignment = WD_TABLE_ALIGNMENT.RIGHT
    set_table_geometry(approval, [720, 980, 560, 1010, 1010], 4280)
    set_table_borders(approval, size=8)

    fill_cell(approval.cell(0, 0), "선결", size=8.3, bold=True)
    fill_cell(approval.cell(0, 1), "")
    fill_cell(approval.cell(0, 2), "지시", size=8.3, bold=True)
    merged_top = approval.cell(0, 3).merge(approval.cell(0, 4))
    fill_cell(merged_top, "")

    reception = approval.cell(1, 0).merge(approval.cell(4, 0))
    fill_cell(reception, "접\n수", size=8.2, bold=True)
    circulation = approval.cell(1, 2).merge(approval.cell(4, 2))
    fill_cell(circulation, "결\n재\n·\n공\n람", size=7.4, bold=True)

    for row_index, label in enumerate(["일자", "접수번호", "처리과", "담당자"], 1):
        fill_cell(approval.cell(row_index, 1), label, size=7.8, bold=True)
        fill_cell(approval.cell(row_index, 3), "")
        fill_cell(approval.cell(row_index, 4), "")

    # Remove the mandatory trailing empty paragraph after the nested table.
    for paragraph in list(right.paragraphs):
        if not paragraph.text:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 0.5

    subject = doc.add_paragraph()
    subject.paragraph_format.space_before = Pt(10)
    subject.paragraph_format.space_after = Pt(7)
    subject.paragraph_format.keep_with_next = True
    label = subject.add_run("제    목 : ")
    set_run_font(label, SERIF, 10.8, True)
    value = subject.add_run("내부 네트워크 현황 확인을 위한 방화벽 조회용 계정 요청")
    set_run_font(value, SERIF, 10.8, True)
    set_paragraph_border(subject, bottom={"size": 8, "space": 2, "color": BLACK})

    add_body_paragraph(doc, "귀사의 일익 번창하심을 기원합니다.")
    add_body_paragraph(
        doc,
        "당사 내부 네트워크 현황 확인 및 통신 상태 점검을 위하여 아래와 같이 "
        "KCTV 방화벽의 조회용(Read-Only) 계정 생성을 요청드리오니 협조하여 "
        "주시기 바랍니다.",
    )

    below = doc.add_paragraph()
    below.alignment = WD_ALIGN_PARAGRAPH.CENTER
    below.paragraph_format.space_before = Pt(2)
    below.paragraph_format.space_after = Pt(5)
    run = below.add_run("-   아      래   -")
    set_run_font(run, SERIF, 10.3, True)

    info = doc.add_table(rows=3, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(info, [2200, 7440], 9640)
    set_table_borders(info, color=BLACK, size=7)
    rows = [
        ("요청 목적", "내부 네트워크 현황 확인 및 통신 점검"),
        ("요청 권한", "조회 전용 (Read-Only)"),
        ("신청 계정 정보", "이름/부서: 최훈철 / 전산팀\n연락처: 010-8660-5032"),
    ]
    for row, (label_text, value_text) in zip(info.rows, rows):
        shade_cell(row.cells[0], LIGHT_GRAY)
        fill_cell(row.cells[0], label_text, size=9.1, bold=True, font=SANS)
        fill_cell(
            row.cells[1],
            value_text,
            size=9.2,
            bold=False,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font=SERIF,
        )
        set_cell_margins(row.cells[0], top=105, bottom=105, start=100, end=100)
        set_cell_margins(row.cells[1], top=105, bottom=105, start=150, end=120)

    closing = add_body_paragraph(doc, "감사합니다.  끝.")
    closing.paragraph_format.space_before = Pt(8)
    closing.paragraph_format.space_after = Pt(0)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(108)
    spacer_run = spacer.add_run("\u00A0")
    set_run_font(spacer_run, SERIF, 1)

    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signature.paragraph_format.space_before = Pt(0)
    signature.paragraph_format.space_after = Pt(0)
    signature.paragraph_format.keep_together = True
    run = signature.add_run(
        "㈜타미우스골프앤빌리지  대표이사  김  양  옥"
    )
    set_run_font(run, SANS, 13.2, True)

    settings = doc.settings._element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    build(Path(sys.argv[1]))
