from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def block_items(parent):
    parent_element = parent.element.body
    for child in parent_element.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def paragraph_data(paragraph: Paragraph):
    return {
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style else None,
        "alignment": str(paragraph.alignment),
        "runs": [
            {
                "text": run.text,
                "font": run.font.name,
                "size_pt": run.font.size.pt if run.font.size else None,
                "bold": run.bold,
            }
            for run in paragraph.runs
        ],
    }


def main() -> None:
    document = Document(Path(sys.argv[1]))
    output = []
    for item in block_items(document):
        if isinstance(item, Paragraph):
            output.append({"kind": "paragraph", **paragraph_data(item)})
        else:
            output.append(
                {
                    "kind": "table",
                    "rows": [
                        [
                            [paragraph_data(p) for p in cell.paragraphs]
                            for cell in row.cells
                        ]
                        for row in item.rows
                    ],
                }
            )
    result = {
        "sections": len(document.sections),
        "blocks": output,
        "headers": [
            [paragraph_data(p) for p in section.header.paragraphs]
            for section in document.sections
        ],
        "footers": [
            [paragraph_data(p) for p in section.footer.paragraphs]
            for section in document.sections
        ],
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
